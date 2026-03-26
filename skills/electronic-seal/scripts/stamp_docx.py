#!/usr/bin/env python3
"""
Stamp a seal image onto a Word (.docx) document at specified locations.

Usage:
    python3 stamp_docx.py input.docx seal.png --output output.docx --markers "授权人（盖章）" "被授权人（盖章）"

Options:
    --output, -o     Output docx path (default: input_stamped.docx)
    --markers, -m    Text markers in the document where the seal should be inserted
    --width, -w      Seal width in inches (default: 1.5)

Requires: python-docx (pip install python-docx)
"""

import argparse
import os
import sys

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def stamp_document(input_path, seal_path, output_path=None, markers=None, width=1.5):
    """Insert seal image into a docx at paragraphs containing marker text."""
    if not os.path.exists(input_path):
        print(f"Error: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)
    if not os.path.exists(seal_path):
        print(f"Error: Seal image not found: {seal_path}", file=sys.stderr)
        sys.exit(1)

    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_stamped{ext}"

    if markers is None:
        markers = ["盖章"]

    doc = Document(input_path)
    stamped_count = 0

    for p in doc.paragraphs:
        for marker in markers:
            if marker in p.text:
                run = p.add_run()
                run.add_picture(seal_path, width=Inches(width))
                stamped_count += 1
                break  # Only stamp once per paragraph

    doc.save(output_path)
    print(f"Stamped {stamped_count} location(s). Saved: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Stamp a seal onto a Word document")
    parser.add_argument("input", help="Input .docx file path")
    parser.add_argument("seal", help="Seal PNG image path")
    parser.add_argument("-o", "--output", default=None, help="Output .docx path")
    parser.add_argument("-m", "--markers", nargs="+", default=["盖章"],
                        help="Text markers where seal should be placed")
    parser.add_argument("-w", "--width", type=float, default=1.5,
                        help="Seal width in inches (default: 1.5)")
    args = parser.parse_args()

    stamp_document(args.input, args.seal, args.output, args.markers, args.width)


if __name__ == "__main__":
    main()
